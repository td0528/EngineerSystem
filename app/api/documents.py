"""
Documents API router - 资料管理
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_
from typing import Optional, List
from pydantic import BaseModel
from datetime import datetime
import os
import shutil

# Document parsing
try:
    from PyPDF2 import PdfReader
    from docx import Document as DocxDocument
    from openpyxl import load_workbook
except ImportError:
    pass

from app.database.database import get_db
from app.models.documents import Document, DocumentVersion
from app.api.auth import get_current_user

router = APIRouter()

UPLOAD_DIR = "uploads/documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class DocumentCreate(BaseModel):
    title: str
    category: Optional[str] = None
    subcategory: Optional[str] = None
    tags: Optional[str] = None


class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    category: Optional[str] = None
    subcategory: Optional[str] = None
    tags: Optional[str] = None


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text[:10000]  # Limit text length
    except Exception:
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text[:10000]
    except Exception:
        return ""


def extract_text_from_excel(file_path: str) -> str:
    """Extract text from Excel"""
    try:
        wb = load_workbook(file_path, data_only=True)
        text = ""
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell])
                text += row_text + "\n"
        return text[:10000]
    except Exception:
        return ""


@router.get("/")
async def get_documents(
    search: Optional[str] = None,
    category: Optional[str] = None,
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db)
):
    """获取文档列表，支持全文搜索"""
    query = db.query(Document).filter(Document.is_archived == False)
    
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                Document.title.ilike(search_term),
                Document.search_content.ilike(search_term),
                Document.tags.ilike(search_term)
            )
        )
    
    if category:
        query = query.filter(Document.category == category)
    
    total = query.count()
    documents = query.order_by(Document.updated_at.desc()).offset(skip).limit(limit).all()
    
    return {"total": total, "items": documents}


@router.get("/search")
async def search_documents(
    q: str = Query(..., min_length=1),
    limit: int = 20,
    db: Session = Depends(get_db)
):
    """文档全文搜索"""
    search_term = f"%{q}%"
    documents = db.query(Document).filter(
        Document.is_archived == False,
        or_(
            Document.title.ilike(search_term),
            Document.search_content.ilike(search_term)
        )
    ).limit(limit).all()
    
    return documents


@router.get("/{document_id}")
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """获取文档详情"""
    document = db.query(Document).options(
        joinedload(Document.versions)
    ).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    return document


@router.post("/")
async def create_document(
    title: str,
    category: Optional[str] = None,
    subcategory: Optional[str] = None,
    tags: Optional[str] = None,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """上传新文档"""
    # Determine file type
    file_ext = file.filename.split(".")[-1].lower()
    if file_ext not in ["pdf", "docx", "xlsx", "doc", "xls"]:
        raise HTTPException(status_code=400, detail="只支持 PDF, Word, Excel 文件")
    
    # Save file
    file_name = f"doc_{datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, file_name)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Extract text content for search
    search_content = ""
    if file_ext == "pdf":
        search_content = extract_text_from_pdf(file_path)
    elif file_ext in ["docx", "doc"]:
        search_content = extract_text_from_docx(file_path)
    elif file_ext in ["xlsx", "xls"]:
        search_content = extract_text_from_excel(file_path)
    
    # Create document
    document = Document(
        title=title,
        category=category,
        subcategory=subcategory,
        tags=tags,
        current_version="1.0",
        current_file_path=f"/uploads/documents/{file_name}",
        file_type=file_ext,
        file_size=os.path.getsize(file_path),
        search_content=search_content,
        uploaded_by=1  # TODO: Get from current user
    )
    db.add(document)
    db.flush()
    
    # Create first version
    version = DocumentVersion(
        document_id=document.id,
        version="1.0",
        file_path=document.current_file_path,
        file_size=document.file_size,
        change_notes="初始版本",
        uploaded_by=1
    )
    db.add(version)
    
    db.commit()
    db.refresh(document)
    
    return document


@router.post("/{document_id}/versions")
async def upload_new_version(
    document_id: int,
    version: str,
    change_notes: Optional[str] = None,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """上传新版本"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # Save file
    file_ext = file.filename.split(".")[-1].lower()
    file_name = f"doc_{document_id}_v{version}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, file_name)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Extract text content
    search_content = ""
    if file_ext == "pdf":
        search_content = extract_text_from_pdf(file_path)
    elif file_ext in ["docx", "doc"]:
        search_content = extract_text_from_docx(file_path)
    elif file_ext in ["xlsx", "xls"]:
        search_content = extract_text_from_excel(file_path)
    
    # Create version record
    doc_version = DocumentVersion(
        document_id=document_id,
        version=version,
        file_path=f"/uploads/documents/{file_name}",
        file_size=os.path.getsize(file_path),
        change_notes=change_notes,
        uploaded_by=1
    )
    db.add(doc_version)
    
    # Update document
    document.current_version = version
    document.current_file_path = doc_version.file_path
    document.file_size = doc_version.file_size
    document.search_content = search_content
    
    db.commit()
    db.refresh(document)
    
    return document


@router.get("/{document_id}/versions")
async def get_document_versions(document_id: int, db: Session = Depends(get_db)):
    """获取文档版本历史"""
    versions = db.query(DocumentVersion).filter(
        DocumentVersion.document_id == document_id
    ).order_by(DocumentVersion.created_at.desc()).all()
    
    return versions


@router.get("/categories/all")
async def get_document_categories(db: Session = Depends(get_db)):
    """获取所有分类"""
    categories = db.query(Document.category).distinct().filter(
        Document.category.isnot(None)
    ).all()
    
    return [c[0] for c in categories]
